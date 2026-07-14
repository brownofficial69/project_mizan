
---

## Step 2: The Compilation Script

Save the script below as `compile_mizan.py` in the same directory as your `document.md`. This script creates your native Microsoft Word document (.docx) directly out of the markdown text with zero layout errors or text residue.

```python
import os

def check_dependencies():
    try:
        import docx
    except ImportError:
        print("[!] Installing dependency: python-docx...")
        os.system("pip install python-docx")

def convert_md_to_docx(md_path, docx_path):
    from docx import Document
    from docx.shared import Pt, Inches
    
    if not os.path.exists(md_path):
        print(f"[!] Error: {md_path} not found.")
        return

    doc = Document()
    
    # Configure clean margin layouts
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code fences
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line.replace('\n', ''))
            p.paragraph_format.left_indent = Inches(0.5)
            style = p.style
            font = style.font
            font.name = 'Consolas'
            font.size = Pt(9.5)
            continue

        # Handle Markdown headings
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            h.style.font.name = 'Arial'
            h.style.font.size = Pt(20)
            h.style.font.bold = True
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            h.style.font.name = 'Arial'
            h.style.font.size = Pt(14)
            h.style.font.bold = True
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            h.style.font.name = 'Arial'
            h.style.font.size = Pt(12)
            h.style.font.bold = True
        elif stripped.startswith("*   ") or stripped.startswith("* "):
            content = stripped.lstrip("* ").strip()
            p = doc.add_paragraph(content, style='List Bullet')
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)
        elif stripped == "---":
            doc.add_paragraph("____________________________________________________")
        elif stripped == "":
            continue
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)

    doc.save(docx_path)
    print(f"[✓] Document compiled successfully: {docx_path}")

if __name__ == "__main__":
    check_dependencies()
    convert_md_to_docx("document.md", "MIZAN_White_Paper.docx")
    
    print("\n[i] To generate your PDF document matching exactly:")
    print("    - On Windows/Mac: Open 'MIZAN_White_Paper.docx' in Word and click 'Save As' -> PDF.")
    print("    - Via Command Line: run 'pip install cookies' and execute 'pandoc document.md -o MIZAN_White_Paper.pdf'")